import streamlit as st
from docx import Document
import openai

def generate_section_content(api_key, idea, section_name):
    """Generate detailed content for a given section of the business plan using the gpt-3.5-turbo model."""
    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user", "content": f"Provide a detailed {section_name.lower()} for a business plan based on the idea: {idea}"}
    ]
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
        api_key=api_key
    )
    return response.choices[0].message['content'].strip()

st.title("AI-Powered Business Plan Generator")

# Sidebar for API key input
st.sidebar.header("Settings")
api_key = st.sidebar.text_input("Enter your OpenAI API Key:", type="password")

# User inputs
business_idea = st.text_input("Describe your business idea:", "")
location = st.text_input("Enter the target location for your business:", "")

if st.button("Generate Business Plan"):
    if not api_key:
        st.warning("Please enter your OpenAI API key in the sidebar.")
    else:
        sections = [
            "Executive Summary",
            "Company Description",
            "Market Analysis",
            "Organization & Management",
            "Service or Product Line",
            "Marketing & Sales",
            "Funding Request",
            "Financial Projections"
        ]

        # Create a .docx document
        doc = Document()
        doc.add_heading('Business Plan', 0)

        for section in sections:
            content = generate_section_content(api_key, business_idea, section)
            doc.add_heading(section, level=1)
            doc.add_paragraph(content)

        doc_filename = "/tmp/business_plan.docx"
        doc.save(doc_filename)

        # Display business plan sections in Streamlit
        for section in sections:
            st.write(f"## {section}")
            st.write(generate_section_content(api_key, business_idea, section))

        # Provide download link in Streamlit
        st.markdown(f"[Download Business Plan](file://{doc_filename})")
